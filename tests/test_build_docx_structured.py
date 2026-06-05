from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from build_docx_structured import (
    clean_paragraph,
    detect_repeated_patterns,
    join_hyphenated_lines,
    normalize_quotes,
    normalize_whitespace,
)


def test_detect_repeated_patterns_flags_high_frequency_text() -> None:
    texts = ["Daemon Trevas"] * 8 + ["Texto normal"] * 2 + ["Outro texto"] * 2
    result = detect_repeated_patterns(texts, threshold=0.40)
    assert "Daemon Trevas" in result
    assert "Texto normal" not in result


def test_detect_repeated_patterns_ignores_long_texts() -> None:
    long_text = "Este é um parágrafo longo que não deve ser considerado cabeçalho de página repetido." * 2
    texts = [long_text] * 10 + ["curto"] * 2
    result = detect_repeated_patterns(texts, threshold=0.40)
    assert long_text not in result


def test_detect_repeated_patterns_returns_empty_for_unique_texts() -> None:
    texts = [f"Texto único {i}" for i in range(20)]
    result = detect_repeated_patterns(texts, threshold=0.40)
    assert result == set()


def test_clean_paragraph_discards_repeated_patterns() -> None:
    repeated = {"© Daemon Editora"}
    assert clean_paragraph("© Daemon Editora", repeated) is None


def test_clean_paragraph_discards_lone_numbers() -> None:
    repeated: set[str] = set()
    assert clean_paragraph("42", repeated) is None
    assert clean_paragraph("3", repeated) is None


def test_clean_paragraph_discards_only_symbols() -> None:
    repeated: set[str] = set()
    assert clean_paragraph("•••", repeated) is None
    assert clean_paragraph("___", repeated) is None


def test_clean_paragraph_preserves_normal_text() -> None:
    repeated: set[str] = set()
    result = clean_paragraph("O exobiólogo estuda vida alienígena.", repeated)
    assert result == "O exobiólogo estuda vida alienígena."


def test_join_hyphenated_lines_joins_word_split_across_lines() -> None:
    text = "O persona-\ngem avança"
    assert join_hyphenated_lines(text) == "O personagem avança"


def test_join_hyphenated_lines_does_not_join_before_uppercase() -> None:
    text = "Veja o capí-\nTULO seguinte"
    result = join_hyphenated_lines(text)
    assert "capí-\nTULO" in result


def test_normalize_whitespace_collapses_multiple_spaces() -> None:
    assert normalize_whitespace("texto  com   espaços") == "texto com espaços"


def test_normalize_whitespace_removes_control_characters() -> None:
    assert normalize_whitespace("texto\x00com\x01controles") == "texto com controles"


def test_normalize_quotes_replaces_typographic_quotes() -> None:
    # Unicode left double quote (U+201C), right double quote (U+201D)
    # Unicode left single quote (U+2018), right single quote (U+2019)
    input_text = chr(0x201C) + "Daemon" + chr(0x201D) + " é " + chr(0x2018) + "ótimo" + chr(0x2019)
    expected = '"Daemon" é \'ótimo\''
    result = normalize_quotes(input_text)
    assert result == expected


from docx import Document as DocxDocument
from build_docx_structured import classify_heading_level


def _make_para(tmp_path: Path, text: str, style: str = "Normal", bold: bool = False):
    """Helper: cria um parágrafo python-docx em memória."""
    doc = DocxDocument()
    p = doc.add_paragraph(text)
    try:
        p.style = doc.styles[style]
    except KeyError:
        pass
    if bold and p.runs:
        p.runs[0].bold = True
    return p


def test_classify_heading_level_detects_heading1_by_style(tmp_path: Path) -> None:
    p = _make_para(tmp_path, "Capítulo 1: Introdução", style="Heading 1")
    assert classify_heading_level(p) == 1


def test_classify_heading_level_detects_heading2_by_style(tmp_path: Path) -> None:
    p = _make_para(tmp_path, "Combate corpo a corpo", style="Heading 2")
    assert classify_heading_level(p) == 2


def test_classify_heading_level_detects_chapter_by_pattern(tmp_path: Path) -> None:
    p = _make_para(tmp_path, "Capítulo 3: Magia")
    assert classify_heading_level(p) == 1


def test_classify_heading_level_detects_parte_by_pattern(tmp_path: Path) -> None:
    p = _make_para(tmp_path, "PARTE II — Regras Avançadas")
    assert classify_heading_level(p) == 1


def test_classify_heading_level_returns_none_for_normal_text(tmp_path: Path) -> None:
    p = _make_para(tmp_path, "O personagem pode gastar pontos de aprimoramento.")
    assert classify_heading_level(p) is None


def test_classify_heading_level_returns_none_for_long_bold_text(tmp_path: Path) -> None:
    long = "Este é um texto muito longo que mesmo sendo negrito não deve ser considerado heading pois ultrapassa o limite de caracteres estabelecido."
    p = _make_para(tmp_path, long, bold=True)
    assert classify_heading_level(p) is None


from build_docx_structured import build_structured


def _make_structured_docx(tmp_path: Path) -> Path:
    """Cria DOCX de teste com estrutura mínima."""
    doc = DocxDocument()
    doc.add_paragraph("Texto de capa e copyright")
    doc.add_heading("Capítulo 1: Introdução", level=1)
    doc.add_heading("O Sistema Daemon", level=2)
    doc.add_paragraph("O sistema Daemon usa dados de dez faces para resolver ações.")
    path = tmp_path / "test_book.docx"
    doc.save(str(path))
    return path


def test_build_structured_produces_valid_schema(tmp_path: Path) -> None:
    path = _make_structured_docx(tmp_path)
    result = build_structured(path, source_id="test-book", title="Test Book")
    assert result["version"] == 1
    assert result["source"] == "test-book"
    assert result["sourceType"] == "docx"
    assert isinstance(result["children"], list)
    assert isinstance(result["warnings"], list)


def test_build_structured_creates_frontmatter_before_first_chapter(tmp_path: Path) -> None:
    path = _make_structured_docx(tmp_path)
    result = build_structured(path, source_id="test-book", title="Test Book")
    types = [child["type"] for child in result["children"]]
    assert types[0] == "frontmatter"


def test_build_structured_creates_chapter_node(tmp_path: Path) -> None:
    path = _make_structured_docx(tmp_path)
    result = build_structured(path, source_id="test-book", title="Test Book")
    chapters = [c for c in result["children"] if c["type"] == "chapter"]
    assert len(chapters) == 1
    assert chapters[0]["heading"] == "Capítulo 1: Introdução"
    assert chapters[0]["level"] == 1


def test_build_structured_nests_section_inside_chapter(tmp_path: Path) -> None:
    path = _make_structured_docx(tmp_path)
    result = build_structured(path, source_id="test-book", title="Test Book")
    chapters = [c for c in result["children"] if c["type"] == "chapter"]
    sections = [c for c in chapters[0]["children"] if c["type"] == "section"]
    assert len(sections) == 1
    assert sections[0]["heading"] == "O Sistema Daemon"


def test_build_structured_emits_warning_when_no_chapter_detected(tmp_path: Path) -> None:
    doc = DocxDocument()
    doc.add_paragraph("Texto sem nenhum heading de capítulo.")
    path = tmp_path / "no_chapters.docx"
    doc.save(str(path))
    result = build_structured(path, source_id="no-chapters", title="No Chapters")
    assert any("no_chapter_detected" in w for w in result["warnings"])


def test_build_structured_detects_index_by_heading(tmp_path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("Capítulo 1", level=1)
    doc.add_paragraph("Conteúdo do capítulo.")
    doc.add_heading("Índice", level=1)
    doc.add_paragraph("Conteúdo..............1")
    doc.add_paragraph("Outro item..............2")
    path = tmp_path / "with_index.docx"
    doc.save(str(path))
    result = build_structured(path, source_id="with-index", title="With Index")
    types = [c["type"] for c in result["children"]]
    assert "index" in types
