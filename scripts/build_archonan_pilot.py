from __future__ import annotations

import re
from datetime import datetime
from typing import Iterable

from docx import Document

from common import ROOT, slugify, write_json


SOURCE = "archonan"
TITLE = "Archonan"
SOURCE_CANDIDATES = [
    ROOT / "Livros" / "word" / "Archonan_OCR_alta_qualidade.docx",
    ROOT / "Livros" / "word" / "feito" / "Archonan_OCR_alta_qualidade.docx",
]
SOURCE_PATH = next((path for path in SOURCE_CANDIDATES if path.exists()), SOURCE_CANDIDATES[0])
OUT_PATH = ROOT / "data" / "pilot" / f"{SOURCE}.json"
DOCS_OUT_PATH = ROOT / "docs" / "assets" / "data" / "pilot" / f"{SOURCE}.json"


DROP_EXACT = {
    TITLE,
    "Texto extraido por OCR / camada textual, com limpeza de quebras de linha e caracteres indevidos.",
    "Texto extraído por OCR / camada textual, com limpeza de quebras de linha e caracteres indevidos.",
    "Boa diversao.",
    "Boa diversão.",
}

TEXT_FIXES = {
    "consangiíneos": "congêneres",
    "preferéncia": "preferência",
    "Ma certa vez": "Mas certa vez",
    "idéia": "ideia",
    "idéias": "ideias",
    "chagando": "chegando",
    "mais foi impedida": "mas foi impedida",
    "do qualquer coisa": "do que qualquer coisa",
    "tronar-se": "tornar-se",
    "por que só poderia": "porque só poderia",
    "por que o poderio": "porque o poderio",
    "Magos vermelhos": "Magos Vermelhos",
    "virá um verdadeiro": "vira um verdadeiro",
    "1, 67m": "1,67m",
    "1, 85m": "1,85m",
    "1, 70m": "1,70m",
    "1, 73m": "1,73m",
    "PM: O": "PM: 0",
    "assinada por": "assassinada por",
    "apaixona-se": "apaixonasse",
    "conquista-la": "conquistá-la",
    "entrega-la": "entregá-la",
    "entrega-lo": "entregá-lo",
    "intemar": "internar",
    "intemasse": "internasse",
    "intemaria": "internaria",
    "viajem": "viagem",
    "fregúentes": "frequentes",
    "afim de": "a fim de",
    "mão e não": "mãe e não",
    "Nivel:": "Nível:",
    "Pesquisa/Investigação 459%": "Pesquisa/Investigação 45%",
    "fo passo": "ao passo",
    "Jma vez": "Uma vez",
    "(O)": "O",
    "tomar um Archonanus": "tornar-se um Archonanus",
    "pensando po a aliar 7 grupos": "pensando em se aliar a grupos",
    "Em quanto estiver": "Enquanto estiver",
    "ivandtrash@bol. com. br": "ivanthrash@bol.com.br",
    "ivanthrash@bol. com. br": "ivanthrash@bol.com.br",
}


def normalize_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = text.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")
    text = text.replace("—", "-").replace("–", "-")
    for old, new in TEXT_FIXES.items():
        text = text.replace(old, new)
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    text = re.sub(r"\(\s+", "(", text)
    text = re.sub(r"\s+\)", ")", text)
    return text


def raw_paragraphs() -> list[str]:
    return [paragraph.text for paragraph in Document(SOURCE_PATH).paragraphs if paragraph.text.strip()]


def doc_paragraphs() -> list[str]:
    return [paragraph.text for paragraph in Document(SOURCE_PATH).paragraphs]


def is_page_noise(text: str) -> bool:
    return bool(re.fullmatch(r"Página \d+", text, flags=re.IGNORECASE))


def should_join(previous: str, current: str) -> bool:
    if not previous or not current:
        return False
    if current.startswith(("Local de Nascimento:", "Atributos:", "Aprimoramentos:", "Caminhos:", "Perícias:", "MAX:", "MIKE:")):
        return False
    if previous.endswith((",", ":", "-", "/", "\\")):
        return True
    if current[:1].islower() and not previous.endswith((".", "!", "?", ":", ";", ")")):
        return True
    last_word = previous.split()[-1].lower().strip(".,;:!?")
    if last_word in {"de", "do", "da", "dos", "das", "em", "por", "com", "para", "que", "o", "os", "as", "um", "uma", "no", "na", "e", "ou", "se", "ao", "à"}:
        return True
    return False


def clean(values: Iterable[str]) -> list[str]:
    paragraphs: list[str] = []
    for raw in values:
        text = normalize_text(raw)
        if not text or text in DROP_EXACT or is_page_noise(text):
            continue
        if paragraphs and should_join(paragraphs[-1], text):
            paragraphs[-1] = normalize_text(f"{paragraphs[-1]} {text}")
        else:
            paragraphs.append(text)
    return paragraphs


def section(section_id: str, title: str, area: str, paragraphs: list[str]) -> dict:
    return {"id": section_id, "title": title, "area": area, "paragraphs": paragraphs}


def typed_item(
    title: str,
    area: str,
    kind: str,
    section_title: str,
    paragraphs: list[str],
    sections: list[dict] | None = None,
) -> dict:
    return {
        "id": slugify(title),
        "title": title,
        "area": area,
        "kind": kind,
        "sectionId": slugify(section_title),
        "sectionTitle": section_title,
        "paragraphs": paragraphs,
        "sections": sections or [section("descricao", "Descrição", area, paragraphs)],
    }


def stat_parts(text: str) -> dict[str, str]:
    marker_pattern = r"(Local de Nascimento:|Atributos:|Aprimoramentos:|Caminhos:|Perícias:)"
    matches = list(re.finditer(marker_pattern, text))
    parts: dict[str, str] = {}
    for index, match in enumerate(matches):
        title = match.group(1).rstrip(":")
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        value = normalize_text(text[start:end])
        if value:
            parts[title] = value
    return parts


def npc_stat_sections(stats: list[str]) -> list[dict]:
    parts: dict[str, str] = {}
    for stat in stats:
        parts.update(stat_parts(stat))

    sections: list[dict] = []
    if parts.get("Atributos"):
        sections.append(section("atributos", "Atributos", "criaturas_npcs", [parts["Atributos"]]))
    if parts.get("Perícias"):
        sections.append(section("pericias-e-combate", "Perícias e Combate", "criaturas_npcs", [parts["Perícias"]]))

    abilities = []
    for key in ("Aprimoramentos", "Caminhos"):
        if parts.get(key):
            abilities.append(f"{key}: {parts[key]}")
    if abilities:
        sections.append(section("habilidades", "Habilidades", "criaturas_npcs", abilities))

    if parts.get("Local de Nascimento"):
        sections.append(section("ficha", "Ficha", "criaturas_npcs", [parts["Local de Nascimento"]]))
    return sections


def npc_item(title: str, history: list[str], personality: list[str], stats: list[str]) -> dict:
    sections: list[dict] = npc_stat_sections(stats)
    if history:
        sections.append(section("historia", "História", "criaturas_npcs", history))
    if personality:
        sections.append(section("personalidade-e-objetivos", "Personalidade e Objetivos", "criaturas_npcs", personality))
    return typed_item(
        title,
        "criaturas_npcs",
        "npc",
        "NPC",
        [paragraph for block in sections for paragraph in block["paragraphs"]],
        sections,
    )


def dual_npc_item(title: str, history: list[str], personality: list[str], max_stats: str, mike_stats: str) -> dict:
    max_parts = stat_parts(max_stats)
    mike_parts = stat_parts(mike_stats)
    sections: list[dict] = []
    if max_parts.get("Atributos") or mike_parts.get("Atributos"):
        values = []
        if max_parts.get("Atributos"):
            values.append(f"Max: {max_parts['Atributos']}")
        if mike_parts.get("Atributos"):
            values.append(f"Mike: {mike_parts['Atributos']}")
        sections.append(section("atributos", "Atributos", "criaturas_npcs", values))
    if max_parts.get("Perícias") or mike_parts.get("Perícias"):
        values = []
        if max_parts.get("Perícias"):
            values.append(f"Max: {max_parts['Perícias']}")
        if mike_parts.get("Perícias"):
            values.append(f"Mike: {mike_parts['Perícias']}")
        sections.append(section("pericias-e-combate", "Perícias e Combate", "criaturas_npcs", values))
    abilities = []
    for label, parts in (("Max", max_parts), ("Mike", mike_parts)):
        for key in ("Aprimoramentos", "Caminhos"):
            if parts.get(key):
                abilities.append(f"{label} - {key}: {parts[key]}")
    if abilities:
        sections.append(section("habilidades", "Habilidades", "criaturas_npcs", abilities))
    ficha = []
    if max_parts.get("Local de Nascimento"):
        ficha.append(f"Max: {max_parts['Local de Nascimento']}")
    if mike_parts.get("Local de Nascimento"):
        ficha.append(f"Mike: {mike_parts['Local de Nascimento']}")
    if ficha:
        sections.append(section("ficha", "Ficha", "criaturas_npcs", ficha))
    if history:
        sections.append(section("historia", "História", "criaturas_npcs", history))
    if personality:
        sections.append(section("personalidade-e-objetivos", "Personalidade e Objetivos", "criaturas_npcs", personality))
    return typed_item(
        title,
        "criaturas_npcs",
        "npc",
        "NPC",
        [paragraph for block in sections for paragraph in block["paragraphs"]],
        sections,
    )


def build_lore_item() -> dict:
    sections = [
        section(
            "ficha-da-sociedade",
            "Ficha da Sociedade",
            "cenarios_lore",
            [
                "Fundação: França, século XIX.",
                "Base: Paris.",
                "Atuação: França, principalmente Espanha, Austrália e América do Sul.",
                "Personalidades: Anjo Blanquiel, Demônio Trauros, Proudhon, Bakunin, Kropotkin, Oscar Wilde, Tolstoi, George Orwell, Aldous Huxley, Picasso, Alex Confort, Herbert Read, Emma Goldman, Trilussa, Tucker, Thoreau, Malesta e George Woodcock.",
            ],
        ),
        section(
            "background",
            "Background",
            "cenarios_lore",
            [
                "Em meados de 1870, a batalha pelo domínio da Terra, liderada por Anjos, Demônios e Magos poderosos, já estava bem estruturada. Com exceção dos Templários e das AGNI, quase não havia sociedades secretas tentando restabelecer a paz e a harmonia na Terra.",
                "Ao ver o sangue derramado na Terra, o Serafim Blanquiel desceu para acabar com a guerra. Ele se aliou a Proudhon, Bakunin e Kropotkin, que já lutavam pelo mesmo objetivo, e juntos formularam conceitos sobre o exílio de entidades sobrenaturais malignas da Terra e a conversão de Magos e entidades benignas à Anarquia.",
                "Blanquiel, Proudhon, Bakunin e Kropotkin fundaram a Sociedade Archonan, do grego Archon, governo, e an, sem. O objetivo era devolver liberdade à humanidade e fazê-la atuar como uma unidade cooperativa.",
                "Depois da fundação, o demônio Trauros, vindo de Arkanun e movido por ideologia semelhante e culpa pelas mortes causadas por seus congêneres, aliou-se aos Archonanus. A Sociedade passou a dominar o Caminho Natural Humano e recebeu de Trauros os Caminhos Arcanos da Terra e da Água.",
                "A Archonan não tem aliados fixos e atua praticamente sozinha, com alianças poucas e temporárias com Templários ou AGNI. Seus inimigos incluem sociedades secretas, governos, organizações hierárquicas, Anjos, Demônios, Magos e Iluminados.",
            ],
        ),
        section(
            "historia-recente",
            "História Recente",
            "cenarios_lore",
            [
                "A Sociedade cresceu consideravelmente, destruiu pequenas ordens com objetivos de conquista, implantou regimes anarquistas em pequenas cidades espanholas e enfrentou agentes Iluminados em países da Europa entre 1894 e 1901.",
                "Por volta de 1905, os Iluminados passaram a acompanhar a Sociedade de perto. Em 1907, infiltraram agentes e atacaram uma sede da Archonan em Paris. A batalha causou baixas nos dois lados; o quinteto fundador expulsou os invasores, mas Kropotkin foi morto por Max Weber, agente dos Iluminados, antes de Weber ser destruído por Trauros.",
                "Hoje, apesar de jovem se comparada a sociedades milenares, a Archonan tem contingente atuante, influência na Austrália e ramificações em boa parte da Europa e da América do Sul.",
            ],
        ),
        section(
            "caracteristicas",
            "Características",
            "cenarios_lore",
            [
                "A Archonan não tem vínculo com o Arkanun Arcanorum, e seus quatro fundadores ainda vivos permanecem ao lado dos membros sob forma humana, pois a Sociedade não possui hierarquia.",
                "Os Archonanus seguem cinco mandamentos: jamais negar ajuda a outro Archonanus; matar apenas em último caso; jamais se submeter a hierarquia; combater entidades sobrenaturais, sociedades secretas ou pessoas com objetivos de domínio terrestre; e pregar o movimento anarquista.",
                "O mandamento de matar apenas em último caso é frequentemente ignorado, por causa do perigo constante ao qual os membros são submetidos. A Sociedade discute a possibilidade de abandoná-lo e de se aliar a grupos terroristas ou anarquistas interessados em seu apoio.",
            ],
        ),
        section(
            "organizacao-e-iniciacao",
            "Organização e Iniciação",
            "cenarios_lore",
            [
                "Os Archonanus se organizam em sedes que funcionam como comunidades, onde integrantes podem viver. Na Austrália e na América do Sul existem pequenas cidades sob regime anarquista, embora muitos membros vivam em cidades comuns para manter a Sociedade informada.",
                "Para tornar-se um Archonanus, a pessoa deve concordar com os objetivos da Sociedade e passar por um ritual simbólico de purificação: sobreviver uma semana em uma cidade desconhecida, apenas com as roupas do corpo, sem dinheiro ou pertences, respeitando as leis locais.",
                "O iniciado é observado por outros Archonanus. Caso roube ou cometa outro ato inválido, sua entrada na Sociedade é recusada. Se completar o ritual, recebe um anel com o símbolo da Archonan e passa a compartilhar conhecimentos com os demais membros.",
            ],
        ),
        section(
            "graus",
            "Graus",
            "cenarios_lore",
            [
                "A Archonan não é dividida em graus de poder. Todos estão no mesmo nível, inclusive os fundadores.",
                "Em vez de graus, a Sociedade é dividida em Grupos de Profissão, cada um especializado em uma profissão ou perícia, como Caminho Arcano da Terra, Caminho Arcano da Água, Caminho Humano, Espiões e Atiradores.",
                "Esses grupos ensinam técnicas da profissão. O estudo varia de um a cinco anos. Após concluir uma formação, o Archonanus recebe anel, adorno ou tatuagem que indica sua aptidão para realizar tarefas naquela profissão. O direito de aprender novas habilidades vem de créditos obtidos por feitos em favor da Sociedade.",
            ],
        ),
        section(
            "numeros",
            "Números",
            "cenarios_lore",
            [
                "O bloco final do texto reúne citações de Trilussa, Rousseau e Thoreau sobre números, liberdade, governo e anarquismo, servindo como tom ideológico da Sociedade Archonan.",
            ],
        ),
    ]
    return typed_item(
        "Cenarios/Lore - Archonan",
        "cenarios_lore",
        "setting",
        "Cenário/Lore",
        [paragraph for block in sections for paragraph in block["paragraphs"]],
        sections,
    )


def build_payload() -> dict:
    raw = doc_paragraphs()

    paulo_history = clean([raw[22], raw[24], raw[25], raw[26], raw[27]])
    paulo_personality = clean([raw[28]])
    paulo_stats = clean([raw[29]])

    richard_history = clean([raw[31], raw[32], raw[34], raw[35], raw[36], raw[37]])
    richard_personality = clean([raw[38]])
    richard_stats = clean([raw[39]])

    veronica_history = clean([raw[42], raw[43], raw[44], raw[45], raw[46]])
    veronica_personality_and_stats = clean([f"{raw[47]} {raw[48]}"])
    veronica_text = veronica_personality_and_stats[0]
    marker = "Local de Nascimento:"
    veronica_personality = [normalize_text(veronica_text.split(marker, 1)[0])]
    veronica_stats = [normalize_text(f"{marker}{veronica_text.split(marker, 1)[1]}")]

    max_history = clean([raw[51], raw[52], raw[53], raw[54]])
    max_text = normalize_text(raw[55])
    max_marker = "MAX: Local de Nascimento:"
    max_personality = [normalize_text(max_text.split(max_marker, 1)[0])]
    max_stats = clean([f"Local de Nascimento:{max_text.split(max_marker, 1)[1]} {raw[57]}"])
    mike_stats = clean([raw[58].replace("MIKE: ", "")])

    sections = [
        build_lore_item(),
        npc_item("Paulo de Oliveira", paulo_history, paulo_personality, paulo_stats),
        npc_item("Richard Stanford", richard_history, richard_personality, richard_stats),
        npc_item("Verônica Galdini", veronica_history, veronica_personality, veronica_stats),
        dual_npc_item("Max/Mike Sterenlicht", max_history, max_personality, max_stats[0], mike_stats[0]),
    ]

    return {
        "version": 1,
        "source": SOURCE,
        "title": TITLE,
        "sourceFile": SOURCE_PATH.name,
        "status": "pilot_review",
        "summary": "Sociedade secreta anarquista Archonan e NPCs Archonanus associados.",
        "areas": ["cenarios_lore", "criaturas_npcs"],
        "groups": [],
        "sections": sections,
        "counts": {
            "cenarios_lore": 1,
            "criaturas_npcs": 4,
            "itens": len(sections),
        },
        "reviewNotes": [
            "Texto revisado antes da catalogação, com correções de OCR e remoção de ruído de página.",
            "As páginas iniciais tinham coluna dupla misturada; o conteúdo foi reorganizado em blocos de sociedade, histórico, características, organização e graus.",
            "Personagens foram catalogados como NPCs, com história, personalidade/objetivos e ficha técnica separados.",
        ],
        "generatedAt": datetime.now().isoformat(timespec="seconds"),
    }


def main() -> None:
    payload = build_payload()
    write_json(OUT_PATH, payload)
    write_json(DOCS_OUT_PATH, payload)
    print(f"Wrote {OUT_PATH}")
    print(f"Wrote {DOCS_OUT_PATH}")
    print(f"Sections: {len(payload['sections'])}; counts: {payload['counts']}")


if __name__ == "__main__":
    main()
